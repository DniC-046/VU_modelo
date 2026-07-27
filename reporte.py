import docx
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    """Aplica color de fondo a una celda de tabla en Word"""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Aplica márgenes internos (padding) a una celda"""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def crear_reporte_tecnico():
    doc = docx.Document()

    # ---------------------------------------------------------
    # 1. CONFIGURACIÓN DE MÁRGENES (GUÍA UTTEC)
    # ---------------------------------------------------------
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.left_margin = Cm(3.0)

    # ESTILOS NATIVOS
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(12)
    style_normal.font.color.rgb = RGBColor(0, 0, 0)
    style_normal.paragraph_format.line_spacing = 1.5
    style_normal.paragraph_format.space_after = Pt(6)
    style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # FUNCIONES AUXILIARES DE REDACCIÓN Y ESTILO
    def add_chapter_title(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text.upper())
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.bold = True
        return p

    def add_section_title(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = False
        return p

    def add_body_paragraph(text):
        return doc.add_paragraph(text)

    def add_figure_placeholder(fig_num, description):
        """Genera una caja rectangular estándar simulando la posición de la imagen con su pie de figura."""
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        cell.width = Inches(5.5)
        
        # Estilos visuales del recuadro
        set_cell_background(cell, "F2F4F7")
        set_cell_margins(cell, top=200, bottom=200, left=200, right=200)
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[ INSERTAR FIGURA {fig_num} AQUÍ: Cuadro rectangular estándar ]\n({description})")
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Pie de figura (centrado debajo de la caja)
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(4)
        p_cap.paragraph_format.space_after = Pt(12)
        
        run_bold = p_cap.add_run(f"Figura {fig_num} ")
        run_bold.font.name = 'Arial'
        run_bold.font.size = Pt(10)
        run_bold.font.bold = True
        
        run_text = p_cap.add_run(description)
        run_text.font.name = 'Arial'
        run_text.font.size = Pt(10)

    # ---------------------------------------------------------
    # 2. PORTADILLA Y SECCIONES PRELIMINARES VACÍAS
    # ---------------------------------------------------------
    # Portadilla
    p_port = doc.add_paragraph()
    p_port.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_port = p_port.add_run("UNIVERSIDAD TECNOLÓGICA DE TECÁMAC\nDIVISIÓN DE TECNOLOGÍAS DE LA INFORMACIÓN Y COMUNICACIÓN\n\n")
    r_port.bold = True
    r_port.font.size = Pt(14)
    
    r_title = p_port.add_run("IMPLEMENTACIÓN DE MODELO DE ANALÍTICA DE DATOS EN VIRTUAL UTTEC\n\nMEMORIA DE ESTADÍA PROFESIONAL\nREPORTE TÉCNICO\n\n")
    r_title.bold = True
    r_title.font.size = Pt(16)
    
    p_port.add_run("PRESENTA:\nDENISSE ODALYS CORONA GUZMÁN\n\nASESORA DE LA ORGANIZACIÓN: MTRA. GABRIELA JIMÉNEZ ENCISO\nASESORA ACADÉMICA: DRA. JOCABED CARRADA MARISCAL\n\nJULIO 2026")
    doc.add_page_break()

    # Secciones vacías obligatorias por norma
    secciones_vacias = [
        "CARTA DE AUTORIZACIÓN DE DIGITALIZACIÓN DE REPORTE TÉCNICO",
        "AGRADECIMIENTOS Y DEDICATORIAS",
        "RESUMEN",
        "ABSTRACT",
        "INTRODUCCIÓN"
    ]
    for sec in secciones_vacias:
        add_chapter_title(sec)
        doc.add_page_break()

    # ---------------------------------------------------------
    # 3. ÍNDICE GENERAL DEPURADO (SIN ENTRADAS DE FIGURAS)
    # ---------------------------------------------------------
    add_chapter_title("ÍNDICE")
    
    indice_elementos = [
        ("RESUMEN", "1"),
        ("ABSTRACT", "2"),
        ("INTRODUCCIÓN", "3"),
        ("OBJETIVOS", "4"),
        ("PROGRAMA Y CRONOGRAMA", "5"),
        ("MARCO TEÓRICO", "6"),
        ("METODOLOGÍA", "12"),
        ("CAPÍTULO 1. ANÁLISIS", "15"),
        ("  1.1.0 Variedad de herramientas identificadas en virtual UTTEC", "16"),
        ("  1.2.0 Herramientas nativas de Moodle implementadas en virtual UTTEC", "17"),
        ("    1.2.1 Tablero de calendario", "17"),
        ("    1.2.3 Bitácora de visualización de actividad de los usuarios", "19"),
        ("    1.2.4 Registro detallado del seguimiento del usuario", "21"),
        ("    1.2.5 Clics de los usuarios", "23"),
        ("  1.3.0 Organización de cursos", "25"),
        ("    1.3.1 Desempeño de los usuarios en los cursos", "27"),
        ("    1.3.2 Calificaciones de los alumnos", "27"),
        ("  1.4.0 Termino de las herramientas nativas Moodle en virtual UTTEC", "29"),
        ("    1.4.1 Recolección de la información analizada", "29"),
        ("CAPÍTULO 2. INSTALACIÓN Y CONFIGURACIÓN DE HERRAMIENTAS PARA EL DESARROLLO DEL MODELO", "31"),
        ("  2.1.0 Instalación", "31"),
        ("    2.1.1 Instalación de Moodle", "31"),
        ("    2.1.2 Instalación de Python", "33"),
        ("    2.1.3 Instalar el entorno virtual de Python", "33"),
        ("    2.1.4 Instalación de Pandas", "34"),
        ("    2.1.5 Instalación de Apache", "35"),
        ("    2.1.6 Instalación de PHP", "37"),
        ("    2.1.7 Instalación de MariaDB", "37"),
        ("    2.1.8 Verificación de existencia de Git", "39"),
        ("    2.1.9 Instalación de Dash de Plotly", "39"),
        ("  2.2.0 Configuración", "40"),
        ("    2.2.1 Configuración de Render", "40"),
        ("    2.2.2 Configuración de GitHub", "41"),
        ("  2.3.0 Información final obtenida", "41"),
        ("CAPÍTULO 3. DISEÑO DE LA APLICACIÓN WEB", "42"),
        ("  3.1 Introducción al Diseño del Sistema", "42"),
        ("  3.2 Diseño de Interfaz e Iconografía", "43"),
        ("    3.2.1 Menú lateral izquierdo", "43"),
        ("    3.2.2 Perfil de usuarios", "44"),
        ("    3.2.3 Alternancia de tema (Modo Claro / Oscuro)", "45"),
        ("  3.3 Visualización de Datos", "46"),
        ("    3.3.1 Gráficas de datos", "46"),
        ("  3.4 Módulos de Usuario", "47"),
        ("    3.4.1 Diseño del apartado de estudiantes", "47"),
        ("    3.4.2 Módulo de diagnóstico pedagógico y predicción por IA", "48"),
        ("  3.5 Consideraciones Generales de Diseño", "49"),
        ("    3.5.1 Diseño general y responsividad", "49"),
        ("    3.5.2 Adaptabilidad y maquetación por tarjetas", "50"),
        ("CAPÍTULO 4. TESTEOS Y FUNCIONALIDAD DEL MODELO", "51"),
        ("CONCLUSIONES", "52"),
        ("REFERENCIAS", "54")
    ]

    for item, pag in indice_elementos:
        p_ind = doc.add_paragraph()
        p_ind.paragraph_format.space_after = Pt(2)
        p_ind.paragraph_format.line_spacing = 1.15
        
        # Puntos de relleno
        run_item = p_ind.add_run(item)
        if item.startswith("CAPÍTULO") or item in ["RESUMEN", "ABSTRACT", "INTRODUCCIÓN", "OBJETIVOS", "MARCO TEÓRICO", "METODOLOGÍA", "CONCLUSIONES", "REFERENCIAS"]:
            run_item.bold = True
            
        dots_len = max(5, 85 - len(item))
        run_dots = p_ind.add_run(" " + "." * dots_len + " " + pag)
        run_dots.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_page_break()

    # ---------------------------------------------------------
    # 4. CAPÍTULOS ANTERIORES (SINTETIZADOS PARA MANTENER ESTRUCTURA)
    # ---------------------------------------------------------
    add_chapter_title("OBJETIVOS")
    add_body_paragraph("Objetivo General: Implementar modelos de analítica de datos en la plataforma virtual UTTEC para optimizar la toma de decisiones pedagógicas y el seguimiento del rendimiento académico.")
    
    add_chapter_title("MARCO TEÓRICO")
    add_body_paragraph("Fundamentación teórica orientada al uso de Python, Pandas, Dash de Plotly y la metodología CRISP-DM...")

    add_chapter_title("METODOLOGÍA")
    add_body_paragraph("Descripción del ciclo de vida del proyecto bajo la metodología CRISP-DM dividida en sus seis fases estándar...")

    add_chapter_title("CAPÍTULO 1. ANÁLISIS")
    add_section_title("1.1.0 Variedad de herramientas identificadas en virtual UTTEC")
    add_body_paragraph("Análisis de las características técnicas del LMS Moodle y sus repositorios de logs nativos...")

    add_chapter_title("CAPÍTULO 2. INSTALACIÓN Y CONFIGURACIÓN DE HERRAMIENTAS PARA EL DESARROLLO DEL MODELO")
    add_section_title("2.1.0 Instalación")
    add_body_paragraph("Detalle del despliegue del entorno en Fedora 40, configuración del entorno virtual de Python, instalación de Pandas, Apache, PHP, MariaDB y Dash.")
    
    # Marcador final de figura del Capítulo 2 para mantener la secuencia numérica
    add_figure_placeholder("3.3", "Estructura del repositorio remoto en GitHub para el modelo de analítica")

    doc.add_page_break()

    # ---------------------------------------------------------
    # 5. CAPÍTULO 3. DISEÑO DE LA APLICACIÓN WEB (DESARROLLADO)
    # ---------------------------------------------------------
    add_chapter_title("CAPÍTULO 3. DISEÑO DE LA APLICACIÓN WEB")

    add_section_title("3.1 Introducción al Diseño del Sistema")
    add_body_paragraph("El diseño de la aplicación web 'Analítica UTTEC' fue concebido bajo premisas estrictas de usabilidad visual, respuesta interactiva y alineación completa a la identidad institucional de la Universidad Tecnológica de Tecámac. La interfaz tiene como propósito central transformar datos complejos extraídos del Moodle institucional en paneles visuales intuitivos que faciliten la identificación de patrones de rendimiento académico y estudiantes en condición de riesgo.")

    add_section_title("3.2 Diseño de Interfaz e Iconografía")

    add_section_title("3.2.1 Menú lateral izquierdo")
    add_body_paragraph("Para la navegación principal de la plataforma se construyó una barra lateral izquierda (Sidebar) con estilo estático e institucional. Se eliminó por completo el uso de emojis e iconografía informal, sustituyéndola por componentes de simbología lineal y minimalista vectorizados. La configuración visual del menú se apoya en el archivo de hojas de estilo 'assets/style.css', integrando espaciados verticales de 12px x 16px por enlace, garantizando un área de interacción amplia y limpia.")
    
    add_figure_placeholder("3.4", "Diseño institucional del menú lateral izquierdo con iconografía minimalista y paleta de colores UTTEC")

    add_section_title("3.2.2 Perfil de usuarios")
    add_body_paragraph("En la vista detallada por estudiante, el sistema incorpora un algoritmo dinámico que procesa la cadena de texto del nombre completo del alumno para extraer de manera automática sus iniciales (por ejemplo, generando 'KA' para el usuario Kevin Alexis Marquez Lopez). Estas iniciales se renderizan dentro de una tarjeta de perfil estilizada con forma circular y bordes contrastantes en Verde UTTEC (#00A859), proveyendo una identidad visual clara sin requerir la carga pesada de fotografías externas.")

    add_figure_placeholder("3.5", "Renderizado dinámico de avatar circular con iniciales nominales del estudiante")

    add_section_title("3.2.3 Alternancia de tema (Modo Claro / Oscuro)")
    add_body_paragraph("A fin de adaptar la plataforma a distintas condiciones de iluminación y preferencias del usuario, se implementó un conmutador dinámico de tema visual (Theme Toggle) en el extremo inferior del menú lateral. Este módulo alterna la interfaz entre el 'Modo Oscuro' (fondo oscuro con tipografía contrastante) y el 'Modo Claro' (fondo blanco institucional). Al conmutar el estado, el icono cambia dinámicamente entre un sol y una luna, reajustando automáticamente las variables CSS de color de texto, contenedores de métricas y los ejes de contraste en las gráficas de Plotly.")

    add_section_title("3.3 Visualización de Datos")

    add_section_title("3.3.1 Gráficas de datos")
    add_body_paragraph("La sección analítica principal integra dos visualizaciones interactivas desarrolladas con Plotly Express y renderizadas sobre componentes 'dcc.Graph' de Dash:\n1. Gráfico de Pastel (Pie Chart): Representa la proporción porcentual del estatus académico global, dividiendo a los alumnos Aprobados (Verde UTTEC #00A859) de aquellos En Riesgo (Rojo #FF4D4D).\n2. Gráfico de Barras Verticales: Exhibe el rendimiento nominal y las calificaciones finales por estudiante. Para resolver el amontonamiento tipográfico en grupos masivos, se configuró una rotación fija de etiquetas de eje a 90°, permitiendo una lectura ordenada y precisa de cada registro.")

    add_figure_placeholder("3.6", "Distribución de gráficos interactivos de Plotly Express para el análisis del grupo")

    add_section_title("3.4 Módulos de Usuario")

    add_section_title("3.4.1 Diseño del apartado de estudiantes")
    add_body_paragraph("El apartado individual de estudiantes se activa de forma interactiva al seleccionar una fila de la tabla o utilizar el buscador nominal rápido. Este módulo despliega cuatro métricas fundamentales organizadas horizontalmente: Calificación Acumulada, Promedio del Grupo, Calificación Máxima y Estatus Académico Actual.")

    add_section_title("3.4.2 Módulo de diagnóstico pedagógico y predicción por IA")
    add_body_paragraph("Debajo de los indicadores cuantitativos, se diseñó una tarjeta contenedora de 'Diagnóstico Pedagógico y Predicción por IA'. Este componente procesa los datos del alumno de forma asíncrona y genera un resumen cualitativo del nivel de riesgo de deserción junto con una lista estandarizada de recomendaciones pedagógicas personalizadas.")

    add_figure_placeholder("3.7", "Estructura completa del módulo individual de estudiante y tarjeta de diagnóstico por IA")

    add_section_title("3.5 Consideraciones Generales de Diseño")

    add_section_title("3.5.1 Diseño general y responsividad")
    add_body_paragraph("La maquetación general utiliza una cuadrícula flexible (CSS Grid) estructurada en tarjetas ('cards') independientes con bordes redondeados y sombras tenues. La interfaz es totalmente responsiva, adaptando el ancho de los contenedores a diferentes resoluciones de pantalla y monitores institucionales.")

    add_section_title("3.5.2 Adaptabilidad y maquetación por tarjetas")
    add_body_paragraph("Los selectores superiores ('dcc.Dropdown') permiten filtrar dinámicamente la información por Carrera, Curso Moodle, Grupo Académico y Búsqueda por Estudiante, actualizando los componentes visuales de la pantalla en tiempo real sin requerir recargas completas de la página.")

    doc.add_page_break()

    # ---------------------------------------------------------
    # 6. CAPÍTULO 4. TESTEOS Y FUNCIONALIDAD DEL MODELO
    # ---------------------------------------------------------
    add_chapter_title("CAPÍTULO 4. TESTEOS Y FUNCIONALIDAD DEL MODELO")
    add_body_paragraph("[Sección reservada para el desarrollo de las pruebas de integración, pruebas de latencia en Render, módulos de auditoría en Excel y validación del modelo de datos].")

    # Guardar documento
    nombre_archivo = "Reporte_Tecnico_UTTEC_Actualizado.docx"
    doc.save(nombre_archivo)
    print(f"✅ Documento generado exitosamente como '{nombre_archivo}'")

if __name__ == "__main__":
    crear_reporte_tecnico()
